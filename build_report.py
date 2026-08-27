from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.enum.style import WD_STYLE_TYPE

OUT = 'revue-veille-tech.docx'
BLUE = RGBColor(46,116,181)
NAVY = RGBColor(31,77,120)
MUTED = RGBColor(89,89,89)

def set_cell_shading(cell, fill):
    tcPr = cell._tc.get_or_add_tcPr(); shd = OxmlElement('w:shd'); shd.set(qn('w:fill'), fill); tcPr.append(shd)

def set_cell_width(cell, dxa):
    tcPr = cell._tc.get_or_add_tcPr(); tcW = tcPr.find(qn('w:tcW'))
    if tcW is None: tcW = OxmlElement('w:tcW'); tcPr.append(tcW)
    tcW.set(qn('w:w'), str(dxa)); tcW.set(qn('w:type'), 'dxa')

def table_geometry(table, widths):
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    tblPr = table._tbl.tblPr
    tblW = tblPr.first_child_found_in('w:tblW'); tblW.set(qn('w:w'), '9360'); tblW.set(qn('w:type'), 'dxa')
    ind = OxmlElement('w:tblInd'); ind.set(qn('w:w'),'120'); ind.set(qn('w:type'),'dxa'); tblPr.append(ind)
    for row in table.rows:
        for cell, width in zip(row.cells, widths):
            set_cell_width(cell, width); cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            tcMar = OxmlElement('w:tcMar')
            for side, value in [('top','80'),('bottom','80'),('start','120'),('end','120')]:
                el = OxmlElement('w:'+side); el.set(qn('w:w'), value); el.set(qn('w:type'),'dxa'); tcMar.append(el)
            cell._tc.get_or_add_tcPr().append(tcMar)

def add_text(doc, text, style=None, bold_lead=None):
    p = doc.add_paragraph(style=style)
    if bold_lead and text.startswith(bold_lead):
        r=p.add_run(bold_lead); r.bold=True; p.add_run(text[len(bold_lead):])
    else: p.add_run(text)
    return p

doc=Document(); sec=doc.sections[0]
for side in ('top_margin','bottom_margin','left_margin','right_margin'): setattr(sec,side, Inches(1))
sec.header_distance=Inches(.492); sec.footer_distance=Inches(.492)
normal=doc.styles['Normal']; normal.font.name='Calibri'; normal._element.rPr.rFonts.set(qn('w:ascii'),'Calibri'); normal._element.rPr.rFonts.set(qn('w:hAnsi'),'Calibri'); normal.font.size=Pt(11); normal.paragraph_format.space_after=Pt(6); normal.paragraph_format.line_spacing=1.1
for name,size,color,before,after in [('Heading 1',16,BLUE,16,8),('Heading 2',13,BLUE,12,6),('Heading 3',12,NAVY,8,4)]:
    st=doc.styles[name]; st.font.name='Calibri'; st.font.size=Pt(size); st.font.color.rgb=color; st.paragraph_format.space_before=Pt(before); st.paragraph_format.space_after=Pt(after)
footer=sec.footer.paragraphs[0]; footer.alignment=WD_ALIGN_PARAGRAPH.RIGHT; rr=footer.add_run('Revue de veille technologique | 27 août 2026'); rr.font.size=Pt(8); rr.font.color.rgb=MUTED

p=doc.add_paragraph(); p.paragraph_format.space_after=Pt(2); r=p.add_run('REVUE DE VEILLE TECHNOLOGIQUE'); r.bold=True; r.font.size=Pt(23); r.font.color.rgb=NAVY
p=doc.add_paragraph(); p.paragraph_format.space_after=Pt(16); r=p.add_run('Trois démarches, corrections proposées et modèle cible'); r.font.size=Pt(14); r.font.color.rgb=MUTED
for label,value in [('Date','27 août 2026'),('Périmètre','lecture manuelle ; alertes/agrégation ; radar collaboratif'),('Statut','recommandations génériques, à adapter à la stack et à l’équipe')]:
    p=doc.add_paragraph(); p.paragraph_format.space_after=Pt(2); a=p.add_run(label+' : '); a.bold=True; p.add_run(value)

doc.add_heading('Réponse directe',1)
add_text(doc,"Les trois démarches sont complémentaires. La lecture manuelle apporte le jugement ; l’automatisation apporte la couverture et la rapidité ; un radar partagé transforme les signaux en décisions. La correction prioritaire est une boucle en trois couches : collecte ciblée, revue humaine courte, décision trimestrielle fondée sur des critères et une preuve d’usage.")

doc.add_heading('Comparaison et corrections',1)
rows=[
('1. Lecture manuelle','Contexte, nuance, signaux faibles.','Dépend des personnes ; peu de traçabilité ; couverture inégale.','Thèmes stables, fiche signal (source, impact, confiance, propriétaire), 30 min/semaine, purge des sources peu utiles.'),
('2. Alertes et agrégation','Rapide et répétable pour CVE, versions, dépréciations, licences.','Bruit ; détection partielle ; alerte ≠ décision.','Séparer action immédiate/découverte ; SLA, responsable, règles d’escalade et motifs de clôture.'),
('3. Radar collaboratif','Priorisation, alignement et lien avec les expérimentations.','Peut être lourd ou hors-sol sans preuve terrain.','Comité trimestriel, champion, critères de passage, date de réexamen et expiration des sujets immobiles.')]
table=doc.add_table(rows=1,cols=4); table.style='Table Grid'; table_geometry(table,[1560,2340,2340,3120])
headers=['Démarche','Forces','Limites','Correction concrète']
for c,t in zip(table.rows[0].cells,headers): set_cell_shading(c,'F2F4F7'); p=c.paragraphs[0]; run=p.add_run(t); run.bold=True
for row in rows:
    cells=table.add_row().cells
    for c,t in zip(cells,row): c.paragraphs[0].add_run(t)

doc.add_heading('Pourquoi ces corrections',1)
doc.add_heading('1. Ne pas confondre collecte et évaluation',2)
add_text(doc,"Les outils de sécurité illustrent bien cette distinction : GitHub précise que les alertes de dépendances sont limitées et propose filtrage, attribution et justification de clôture. L’automatisation doit donc créer une file de tri, pas une recommandation implicite.")
doc.add_heading('2. Exiger une preuve avant l’adoption',2)
add_text(doc,"La méthode Technology Radar de Thoughtworks distingue l’exploration de l’usage réel : « Trial » est réservé aux sujets éprouvés dans un contexte de projet. Utiliser la même logique évite de transformer une nouveauté médiatique en standard interne.")
doc.add_heading('3. Fermer la boucle décisionnelle',2)
add_text(doc,"Une veille utile produit une décision, une expérimentation ou un abandon justifié. Pour chaque sujet : problème ciblé, source primaire, impact, coût/risque, propriétaire, échéance et critère de succès.")

doc.add_heading('Modèle cible : cadence de 30 jours',1)
for title,body in [('Semaine 1 — Cadrer','Choisir thèmes, sources primaires, responsables et taxonomie : sécurité, plateformes, langages/frameworks, pratiques, données/IA.'),('Semaine 2 — Automatiser le factuel','Activer les flux adaptés au portefeuille ; définir triage, SLA et archivage.'),('Semaine 3 — Curater','Publier un digest hebdomadaire de 3 à 5 signaux : changement, impact possible et action proposée.'),('Semaine 4 — Décider','Tenir la première revue de radar : conserver, explorer, expérimenter, adopter ou éviter ; inscrire la décision dans le registre d’architecture.')]:
    p=doc.add_paragraph(style='List Number'); r=p.add_run(title+' — '); r.bold=True; p.add_run(body)

doc.add_heading('Indicateurs à suivre',1)
for x in ['Part des alertes critiques triées dans le SLA.','Ratio signal utile / éléments collectés.','Délai entre signal et décision.','Part des décisions avec source primaire et propriétaire.','Résultats des expérimentations : adoption, abandon justifié, report.']:
    doc.add_paragraph(x,style='List Bullet')

doc.add_heading('Limites',1)
add_text(doc,"Les trois démarches réelles n’ont pas été fournies ; cette revue évalue donc trois modèles usuels. Les fréquences, seuils et SLA sont à ajuster à l’exposition sécurité, la taille de l’équipe et le rythme produit.")
doc.add_heading('Sources',1)
sources=[
('Thoughtworks — FAQ Technology Radar','https://www.thoughtworks.com/radar/faq'),
('Thoughtworks — How to create your enterprise technology radar','https://www.thoughtworks.com/insights/blog/technology-strategy/how-to-create-your-enterprise-technology-radar'),
('GitHub Docs — Dependabot alerts','https://docs.github.com/en/code-security/concepts/supply-chain-security/dependabot-alerts'),
('GitHub Docs — Viewing and updating Dependabot alerts','https://docs.github.com/en/code-security/how-tos/manage-security-alerts/manage-dependabot-alerts/view-dependabot-alerts'),
('OWASP — DevSecOps Guideline: Software Composition Analysis','https://owasp.org/www-project-devsecops-guideline/latest/02d-Software-Composition-Analysis'),
('CISA — Reducing the Significant Risk of Known Exploited Vulnerabilities','https://www.cisa.gov/sites/default/files/publications/Reducing_the_Significant_Risk_of_Known_Exploited_Vulnerabilities_20211103.pdf')]
for label,url in sources:
    p=doc.add_paragraph(style='List Bullet'); p.paragraph_format.space_after=Pt(3); r=p.add_run(label+' — '+url); r.font.size=Pt(9); r.font.color.rgb=RGBColor(5,99,193)
doc.save(OUT)
